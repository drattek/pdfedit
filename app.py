"""
===============================================================================
INSTRUCCIONES PERMANENTES PARA IA / GEMINI (DIRECTIVAS DE EDICIÓN QUIRÚRGICA)
===============================================================================
1. PRESERVACIÓN DE ENDPOINTS:
   - Este script contiene MÚLTIPLES endpoints de FastAPI.
   - NUNCA omitas, resumas ni elimines endpoints existentes a menos que se solicite explícitamente.
   - NUNCA uses comentarios de omisión como "# ... resto del código permanece igual ...".

2. EDICIÓN AISLADA:
   - Modifica ÚNICAMENTE las funciones o endpoints indicados en la solicitud.
   - Mantén intactos los parámetros, lógica e importaciones de los demás endpoints.

3. REGLA DE RESPUESTA COMPLETA:
   - Devuelve SIEMPRE el archivo Python completo ejecutable sin cortes ni elisiones.
===============================================================================
"""

from fastapi import FastAPI, Response, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Optional, List
from io import BytesIO
import base64
import os
import time
import re
import unicodedata
import urllib.request
from datetime import datetime, timedelta
import pdfplumber
from playwright.sync_api import sync_playwright
from reportlab.pdfgen import canvas
from reportlab.lib.colors import white, black
from reportlab.pdfbase.pdfmetrics import stringWidth
from PyPDF2 import PdfReader, PdfWriter

# --- COMPRESORES GRÁFICOS DE ALTA DENSIDAD ---
from PIL import Image
from reportlab.lib.utils import ImageReader

# --- LIBRERÍAS DE DETECCIÓN Y EXTRACCIÓN FACIAL ---
import cv2
import fitz  # PyMuPDF
import numpy as np

# --- GENERADOR DE DOCUMENTOS WORD ---
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONTROL DE VERSIONES ---
VERSION = "1.90 - Fusión Completa (Descarga de Archivos de Almacenamiento Local + Endpoints Doosan/PDF)"
print(f"\n{'='*40}")
print(f" INICIANDO SERVICIO VEGUSA - VERSIÓN: {VERSION}")
print(f" MODO: Producción n8n (Integración Completa v1.70 + v1.80)")
print(f"{'='*40}\n")

# =========================================================
# CARGA AUTÓNOMA DE MODELOS FACIALES
# =========================================================
CASCADE_PATH = os.path.join(os.path.dirname(__file__), "haarcascade_frontalface_default.xml")
CASCADE_URL = "https://raw.githubusercontent.com/opencv/opencv/master/data/haarcascades/haarcascade_frontalface_default.xml"

if not os.path.exists(CASCADE_PATH):
    try:
        print(">>> [FACE DETECTOR] Descargando modelo de rostro OpenCV por primera vez...")
        urllib.request.urlretrieve(CASCADE_URL, CASCADE_PATH)
        print(">>> [FACE DETECTOR] ¡Modelo descargado con éxito!")
    except Exception as e:
        print(f">>> [FACE DETECTOR] Error descargando modelo: {e}")

face_cascade = cv2.CascadeClassifier(CASCADE_PATH)

app = FastAPI(title=f"PDF Edit & Doosan Service v{VERSION} — Vegusa Enterprise")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================================================
# CARGA DEL LOGO LOCAL
# =========================================================
LOGO_B64 = ""
LOGO_PATH = os.path.join(os.path.dirname(__file__), "logo_vegusa.png")

try:
    if os.path.exists(LOGO_PATH):
        with open(LOGO_PATH, "rb") as img_file:
            LOGO_B64 = base64.b64encode(img_file.read()).decode('utf-8')
        print(f"\n>>> [LOGO VEGUSA] ¡Éxito! Imagen cargada ({len(LOGO_B64)} caracteres Base64).")
    else:
        print(f"\n>>> [LOGO VEGUSA] Aviso: No se encontró 'logo_vegusa.png' en: {LOGO_PATH}")
except Exception as e:
    print(f"\n>>> [LOGO VEGUSA] Error al cargar la imagen: {str(e)}")


# ---------------------------------------------------------
# UTILIDAD GLOBAL DE LIMPIEZA DE NOMBRES
# ---------------------------------------------------------
def normalizar_nombre(texto: str) -> str:
    if not texto:
        return "Cliente"
    texto_norm = unicodedata.normalize('NFD', texto)
    texto_sin_acentos = ''.join(c for c in texto_norm if unicodedata.category(c) != 'Mn')
    texto_limpio = texto_sin_acentos.strip().replace(" ", "_")
    return re.sub(r'[^a-zA-Z0-9_\-]', '', texto_limpio)


# ---------------------------------------------------------
# MODELOS DE DATOS (Pydantic)
# ---------------------------------------------------------

class ScrapeRequest(BaseModel):
    user: str
    password: str

class DownloadRequest(BaseModel):
    user: str
    password: str
    shipment_ids: List[str]

class OptimizePDFReq(BaseModel):
    file_b64: str
    file_name: Optional[str] = "ine_optimizada.pdf"

class ExtractFaceReq(BaseModel):
    file_b64: str
    file_name: Optional[str] = "documento.pdf"
    angle: float = 0.0

class Rect(BaseModel):
    page: int = Field(0, description="0-based page index")
    x: float
    y: float
    w: float
    h: float

class CoordinateRequest(BaseModel):
    file_b64: str
    target_text: str

class IncotermReq(BaseModel):
    file_b64: str
    incoterm_change: bool = False
    incoterm_text: str = "DAP"
    area: Optional[Rect] = None
    font_size: Optional[float] = 9.0
    leading: Optional[float] = 11.0
    debug_outline: Optional[bool] = False

class BillShipReq(BaseModel):
    file_b64: str
    bill_to_text: Optional[str] = None
    ship_to_text: Optional[str] = None
    bill_to_area: Optional[Rect] = None
    ship_to_area: Optional[Rect] = None
    font_size: Optional[float] = 9.0
    leading: Optional[float] = 11.0
    debug_outline: Optional[bool] = False

class CustomTextOp(BaseModel):
    text: str
    area: Rect
    font_name: Optional[str] = "Helvetica"
    font_size: Optional[float] = 9.0
    leading: Optional[float] = 11.0
    debug_outline: Optional[bool] = False

class CustomBatchReq(BaseModel):
    file_b64: str
    ops: List[CustomTextOp]

class CutRangeReq(BaseModel):
    file_b64: str
    start_page: int
    final_page: int

class CustomPagesReq(BaseModel):
    file_b64: str
    pages: List[int]

class ReferenceParseReq(BaseModel):
    subject: str
    body: Optional[str] = ""

class PDFRequest(BaseModel):
    html: str
    prefijo: Optional[str] = "Identificacion"
    razon_social: Optional[str] = "Cliente"
    agencia_sucursal: Optional[str] = ""

class WordRequest(BaseModel):
    datosExtraidos: dict
    rostro_b64: Optional[str] = None
    remitente_name: Optional[str] = "Cliente"
    prefijo: Optional[str] = "Identificacion"


# ---------------------------------------------------------
# UTILIDADES INTERNAS PDF Y DOOSAN
# ---------------------------------------------------------

def _load_pdf_from_b64(file_b64: str) -> PdfReader:
    raw = base64.b64decode(file_b64)
    return PdfReader(BytesIO(raw))

def _export(writer: PdfWriter) -> bytes:
    out = BytesIO()
    writer.write(out)
    return out.getvalue()

def _make_overlay(page_width: float, page_height: float, draw_ops):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_width, page_height))
    draw_ops(c)
    c.showPage()
    c.save()
    buf.seek(0)
    return buf

def _overlay_rect_with_text(
    reader: PdfReader,
    writer: PdfWriter,
    rect: Rect,
    text: str,
    font_name: str = "Helvetica",
    font_size: float = 9.0,
    leading: float = 11.0,
    debug_outline: bool = False
):
    page_index = max(0, min(rect.page, len(reader.pages) - 1))
    page = reader.pages[page_index]
    media = page.mediabox
    pw, ph = float(media.width), float(media.height)

    def draw_ops(c):
        if debug_outline:
            c.setStrokeColorRGB(1, 0, 0)
            c.setLineWidth(1.2)
            c.rect(rect.x, rect.y, rect.w, rect.h, fill=False, stroke=True)
            return
        c.setFillColor(white)
        c.rect(rect.x, rect.y, rect.w, rect.h, fill=True, stroke=False)
        c.setFillColor(black)
        c.setFont(font_name, font_size)
        x, y = rect.x + 2, rect.y + rect.h - leading
        max_w = rect.w - 4
        
        for line in (text or '').split("\n"):
            words = line.split(" ")
            cur = ""
            for w in words:
                test = (cur + " " + w).strip() if cur else w
                if stringWidth(test, font_name, font_size) <= max_w:
                    cur = test
                else:
                    c.drawString(x, y, cur)
                    y -= leading
                    cur = w
                    if y < rect.y + 2: return
            if y >= rect.y + 2:
                c.drawString(x, y, cur)
                y -= leading

    overlay_reader = PdfReader(_make_overlay(pw, ph, draw_ops))
    page.merge_page(overlay_reader.pages[0])
    for i, p in enumerate(reader.pages):
        writer.add_page(page if i == page_index else p)


class AuthException(Exception):
    pass

def find_frame_with_selector(page, selector):
    for f in page.frames:
        try:
            if f.locator(selector).count() > 0: return f
        except: continue
    return None

def _doosan_navigate_to_results(context, user, password, f_start, f_end):
    page = context.new_page()
    page.set_default_timeout(90000)
    
    print(f">>> [{VERSION}] Accediendo a Bobcat Login...")
    page.goto('https://dealer.bobcat.com/', wait_until="networkidle", timeout=120000)
    page.wait_for_selector('input[name="identifier"]', timeout=60000)
    page.fill('input[name="identifier"]', user)
    pass_field = page.locator('input[name="credentials.passcode"]')
    pass_field.fill(password)
    time.sleep(1)
    pass_field.press("Enter")
    
    time.sleep(3)
    error_selector = ".infobox-error, .okta-form-infobox-error, .o-form-has-errors"
    if page.locator(error_selector).first.is_visible():
        raise AuthException("usuario o contraseña erróneos")

    page.wait_for_load_state('networkidle', timeout=60000)
    page.wait_for_selector('a:has-text("Doobiz")', timeout=45000)
    with context.expect_page(timeout=90000) as new_page_info:
        page.click('a:has-text("Doobiz")', force=True)
    doobiz_page = new_page_info.value
    doobiz_page.wait_for_load_state('networkidle', timeout=90000)
    
    doobiz_page.evaluate("tlnMoveMenu('ROLES://portal_content/cbt/common/roles/parts/com.di.cbt.cbt_parts_dealer/parts_2/status/shipment_status',0,'width=500,height=750','');")
    time.sleep(15) 
    
    content_frame = None
    for _ in range(15):
        content_frame = doobiz_page.frame(name="isolatedWorkArea")
        if content_frame:
            try:
                if content_frame.locator("#fromPeriod").count() > 0: break
            except: pass
        time.sleep(1)

    if not content_frame:
        content_frame = find_frame_with_selector(doobiz_page, '#fromPeriod')

    if not content_frame:
        raise Exception("No se localizó el frame de trabajo")

    content_frame.locator('#fromPeriod').wait_for(state="visible", timeout=60000)
    content_frame.locator('#fromPeriod').fill(f_start)
    content_frame.locator('#toPeriod').fill(f_end)
    content_frame.locator('button:has-text("Search")').click()
    
    time.sleep(25)
    results_frame = find_frame_with_selector(doobiz_page, 'input[type="radio"]') or content_frame
    return doobiz_page, content_frame, results_frame


# ---------------------------------------------------------
# ENDPOINTS DE LA API
# ---------------------------------------------------------

# --- ENDPOINT 1: EXTRAER ROSTRO (Versión 1.70 - Rotación Corregida) ---
@app.post("/extraer_rostro")
async def extraer_rostro(req: ExtractFaceReq):
    try:
        file_bytes = base64.b64decode(req.file_b64)
        file_name = (req.file_name or "documento.pdf").lower()
        img_cv2 = None

        print(f"\n>>> [EXTRACTOR ROSTRO NATIVO] Procesando archivo: {file_name} (Ángulo recibido: {req.angle}°)")

        if file_name.endswith('.pdf'):
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            if len(pdf_doc) == 0:
                return {"status": "error", "message": "El archivo PDF está vacío.", "rostro_b64": None}
            
            page = pdf_doc.load_page(0)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            
            if pix.n == 4:
                img_cv2 = cv2.cvtColor(img_array, cv2.COLOR_RGBA2BGR)
            else:
                img_cv2 = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        else:
            nparr = np.frombuffer(file_bytes, np.uint8)
            img_cv2 = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

        if img_cv2 is None:
            return {"status": "error", "message": "No se pudo decodificar la imagen del archivo.", "rostro_b64": None}

        # --- 1. ENDEREZADO MATEMÁTICO DIRECTO ---
        if req.angle != 0:
            (h, w) = img_cv2.shape[:2]
            center = (w // 2, h // 2)
            
            angulo_corregido = float(req.angle)
            M = cv2.getRotationMatrix2D(center, angulo_corregido, 1.0)
            
            abs_cos = abs(M[0, 0])
            abs_sin = abs(M[0, 1])
            bound_w = int(h * abs_sin + w * abs_cos)
            bound_h = int(h * abs_cos + w * abs_sin)
            
            M[0, 2] += bound_w / 2 - center[0]
            M[1, 2] += bound_h / 2 - center[1]
            
            img_cv2 = cv2.warpAffine(img_cv2, M, (bound_w, bound_h), borderValue=(255, 255, 255))
            print(f">>> [EXTRACTOR ROSTRO] Documento enderezado a {angulo_corregido}° exitosamente.")

        # --- 2. BÚSQUEDA DEL ROSTRO ---
        gray = cv2.cvtColor(img_cv2, cv2.COLOR_BGR2GRAY)
        gray_eq = cv2.equalizeHist(gray)

        faces = face_cascade.detectMultiScale(gray_eq, scaleFactor=1.1, minNeighbors=5, minSize=(60, 60))
            
        if len(faces) == 0:
            faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=4, minSize=(50, 50))

        if len(faces) == 0:
            print(">>> [EXTRACTOR ROSTRO] Aviso: No se localizó ningún rostro en el documento.")
            return {"status": "error", "message": "No se detectó ningún rostro en el documento.", "rostro_b64": None}

        faces = sorted(faces, key=lambda b: b[2] * b[3], reverse=True)
        (x, y, w, h) = faces[0]

        h_img, w_img, _ = img_cv2.shape
        margen_y = int(h * 0.25)
        margen_x = int(w * 0.20)
        
        y1 = max(0, y - margen_y)
        y2 = min(h_img, y + h + margen_y)
        x1 = max(0, x - margen_x)
        x2 = min(w_img, x + w + margen_x)

        rostro_recortado = img_cv2[y1:y2, x1:x2]
        _, buffer = cv2.imencode('.jpg', rostro_recortado, [int(cv2.IMWRITE_JPEG_QUALITY), 92])
        rostro_b64 = base64.b64encode(buffer).decode('utf-8')

        print(">>> [EXTRACTOR ROSTRO] ¡Éxito! Rostro extraído correctamente mediante OpenCV.")

        return {
            "status": "success",
            "message": "Rostro extraído correctamente.",
            "rostro_b64": rostro_b64
        }

    except Exception as e:
        print(f">>> [ERROR EXTRACTOR ROSTRO]: {str(e)}")
        return {"status": "error", "message": str(e), "rostro_b64": None}


# --- ENDPOINT 2: GENERAR PDF (Versión 1.70 - Nombres Dinámicos) ---
@app.post("/generate-pdf", response_class=Response)
def generate_pdf(req: PDFRequest):
    try:
        html_final = req.html
        if LOGO_B64:
            html_final = html_final.replace("{{LOGO_VEGUSA}}", LOGO_B64)

        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-setuid-sandbox", "--disable-dev-shm-usage"]
            )
            context = browser.new_context()
            page = context.new_page()
            page.set_content(html_final)
            page.wait_for_load_state("networkidle")
            
            pdf_bytes = page.pdf(
                format="Letter",
                print_background=True,
                margin={"top": "15mm", "bottom": "15mm", "left": "12mm", "right": "12mm"}
            )
            context.close()
            browser.close()
            
        prefijo_limpio = normalizar_nombre(req.prefijo or "Identificacion")
        razon_limpia = normalizar_nombre(req.razon_social or "Cliente")
        
        if req.agencia_sucursal and req.agencia_sucursal.lower() != "general":
            sucursal_limpia = normalizar_nombre(req.agencia_sucursal)
            filename_final = f"{prefijo_limpio}_{razon_limpia}_{sucursal_limpia}.pdf"
        else:
            filename_final = f"{prefijo_limpio}_{razon_limpia}.pdf"
            
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={filename_final}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        print(f">>> [ERROR GENERATE-PDF]: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error al renderizar el PDF: {str(e)}"
        )


# --- ENDPOINT 3: GENERAR WORD (Versión 1.70 - Nombres Dinámicos) ---
@app.post("/generate-word", response_class=Response)
def generate_word(req: WordRequest):
    try:
        print("\n>>> [GENERATE WORD] Creando archivo Word (.docx)...")
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        if os.path.exists(LOGO_PATH):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.8))

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("REPORTE DE VALIDACIÓN DE IDENTIFICACIÓN (INE)")
        run_title.bold = True
        run_title.font.size = Pt(15)
        run_title.font.color.rgb = RGBColor(15, 23, 42)

        if req.rostro_b64:
            try:
                img_bytes = base64.b64decode(req.rostro_b64)
                img_stream = BytesIO(img_bytes)
                p_face = doc.add_paragraph()
                p_face.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_face.paragraph_format.space_before = Pt(10)
                p_face.paragraph_format.space_after = Pt(10)
                p_face.add_run().add_picture(img_stream, width=Inches(1.2))
            except Exception as e_img:
                print(f"--> Aviso al insertar rostro en Word: {e_img}")

        datos = req.datosExtraidos or {}
        dir_data = datos.get("direccion", {})

        tabla = doc.add_table(rows=0, cols=2)
        tabla.style = 'Table Grid'

        filas = [
            ("Nombre Completo:", f"{datos.get('nombre', '')} {datos.get('apellidos', '')}".strip()),
            ("Fecha de Nacimiento:", datos.get("fechaNacimiento", "N/D")),
            ("Sexo:", datos.get("sexo", "N/D")),
            ("CURP:", datos.get("curp", "N/D")),
            ("Clave de Elector:", datos.get("claveElector", "N/D")),
            ("CIC / IDMEX:", datos.get("cicIdmex", "N/D")),
            ("ID Ciudadano:", datos.get("idCiudadano", "N/D")),
            ("Estado de Vigencia:", f"{'✔️ VIGENTE' if datos.get('vigente') else '❌ NO VIGENTE'} (Hasta {datos.get('añoVigencia', 'N/D')})"),
            ("Calle y Número:", dir_data.get("calleNumero", "N/D")),
            ("Colonia:", dir_data.get("colonia", "N/D")),
            ("Ciudad / Municipio:", dir_data.get("ciudad", "N/D")),
            ("Estado:", dir_data.get("estado", "Guanajuato")),
            ("Código Postal:", dir_data.get("cp", "N/D"))
        ]

        for etiqueta, valor in filas:
            row_cells = tabla.add_row().cells
            p0 = row_cells[0].paragraphs[0]
            r0 = p0.add_run(etiqueta)
            r0.bold = True
            r0.font.size = Pt(10)
            
            p1 = row_cells[1].paragraphs[0]
            r1 = p1.add_run(str(valor))
            r1.font.size = Pt(10)

        p_disc = doc.add_paragraph()
        p_disc.paragraph_format.space_before = Pt(25)
        p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_disc = p_disc.add_run("🔒 Documento para uso interno exclusivo de Grupo Vegusa. Queda estrictamente prohibida la divulgación o difusión de este archivo fuera de la empresa.")
        run_disc.font.size = Pt(8.5)
        run_disc.font.italic = True
        run_disc.font.bold = True
        run_disc.font.color.rgb = RGBColor(220, 38, 38)

        out_buf = BytesIO()
        doc.save(out_buf)
        docx_bytes = out_buf.getvalue()

        nombre_persona_extraido = f"{datos.get('nombre', '')} {datos.get('apellidos', '')}".strip()
        nombre_base = nombre_persona_extraido if nombre_persona_extraido else (req.remitente_name or "Cliente")
        
        nombre_limpio = normalizar_nombre(nombre_base)
        prefijo_limpio = normalizar_nombre(req.prefijo or "Identificacion")
        
        filename_final = f"{prefijo_limpio}_{nombre_limpio}.docx"

        print(f">>> [GENERATE WORD] ¡Éxito! Archivo generado: {filename_final}")

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename_final}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        print(f">>> [ERROR GENERATE-WORD]: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error al generar documento Word: {str(e)}")


# --- ENDPOINT 4: OPTIMIZAR PDF ---
@app.post("/optimizar_pdf", response_class=Response)
async def optimizar_pdf(req: OptimizePDFReq):
    try:
        raw_bytes = base64.b64decode(req.file_b64)
        reader = PdfReader(BytesIO(raw_bytes))
        writer = PdfWriter()

        for page in reader.pages:
            page_optimized = False
            if page.images:
                for image_file_object in page.images:
                    try:
                        image_bytes = image_file_object.data
                        img = Image.open(BytesIO(image_bytes))
                        
                        if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                            background = Image.new("RGB", img.size, (255, 255, 255))
                            if img.mode == "P": img = img.convert("RGBA")
                            background.paste(img, mask=img.split()[-1])
                            img = background
                        else:
                            img = img.convert("RGB")
                        
                        max_pixel_dim = 1200
                        if img.width > max_pixel_dim or img.height > max_pixel_dim:
                            img.thumbnail((max_pixel_dim, max_pixel_dim), Image.Resampling.LANCZOS)
                        
                        compressed_img_buffer = BytesIO()
                        img.save(compressed_img_buffer, format="JPEG", quality=50, optimize=True)
                        compressed_img_buffer.seek(0)
                        
                        page_buffer = BytesIO()
                        canvas_page = canvas.Canvas(page_buffer, pagesize=(612, 792))
                        img_reader = ImageReader(compressed_img_buffer)
                        canvas_page.drawImage(img_reader, 0, 0, width=612, height=792, preserveAspectRatio=True)
                        canvas_page.showPage()
                        canvas_page.save()
                        page_buffer.seek(0)
                        
                        pdf_page_reader = PdfReader(page_buffer)
                        writer.add_page(pdf_page_reader.pages[0])
                        page_optimized = True
                        break  
                    except Exception: continue
            
            if not page_optimized:
                page.scale_to(612, 792)
                writer.add_page(page)

        pdf_data = _export(writer)
        nombre_original = req.file_name or "ine_optimizada.pdf"
        nombre_seguro = normalizar_nombre(nombre_original)
        if not nombre_seguro.endswith(".pdf"): nombre_seguro += ".pdf"

        return Response(
            content=pdf_data,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={nombre_seguro}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al procesar el PDF: {str(e)}")


# =========================================================
# ENDPOINTS RECUPERADOS DE LA VERSIÓN 1.60
# =========================================================

# --- ENDPOINT 5: BUSCA INVOICE DOOSAN ---
@app.post("/buscaInvoice")
def busca_invoice(req: ScrapeRequest):
    today_dt = datetime.now()
    yesterday_dt = today_dt - timedelta(days=1)
    f_start, f_end = yesterday_dt.strftime("%Y.%m.%d"), today_dt.strftime("%Y.%m.%d")
    shipments = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(viewport={'width': 1920, 'height': 1080})
        try:
            _, _, res_frame = _doosan_navigate_to_results(context, req.user, req.password, f_start, f_end)
            rows = res_frame.locator('tr').filter(has=res_frame.locator('input[type="radio"]')).all()
            for row in rows:
                cells = row.locator('td').all()
                if len(cells) > 1:
                    val = cells[1].inner_text().strip().lstrip('0')
                    if val and val != "...": shipments.append(val)
            return {"status": "success", "version": VERSION, "shipments": list(set(shipments))}
        except AuthException as ae:
            return {"status": "auth_error", "user": req.user, "message": str(ae)}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
        finally: browser.close()


# --- ENDPOINT 6: DESCARGA INVOICE DOOSAN ---
@app.post("/descargaInvoice")
def descarga_invoice(req: DownloadRequest):
    today_dt = datetime.now()
    f_start, f_end = (today_dt - timedelta(days=7)).strftime("%Y.%m.%d"), today_dt.strftime("%Y.%m.%d")
    files = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(viewport={'width': 1920, 'height': 1080})
        try:
            db_page, c_frame, res_frame = _doosan_navigate_to_results(context, req.user, req.password, f_start, f_end)
            rows = res_frame.locator('tr').filter(has=res_frame.locator('input[type="radio"]')).all()
            for row in rows:
                cells = row.locator('td').all()
                if len(cells) > 1:
                    ship_id = cells[1].inner_text().strip().lstrip('0')
                    if ship_id in req.shipment_ids:
                        cells[0].click(force=True)
                        time.sleep(2)
                        with db_page.expect_download(timeout=90000) as download_info:
                            c_frame.locator('button:has-text("Commercial Invoice")').click(force=True)
                        download = download_info.value
                        with open(download.path(), "rb") as f:
                            b64 = base64.b64encode(f.read()).decode('utf-8')
                        files.append({"shipment_no": ship_id, "filename": f"{ship_id.zfill(10)}.pdf", "base64": b64})
            return {"status": "success", "files": files}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
        finally: browser.close()


# --- ENDPOINT 7: EXTRACT COORDINATES ---
@app.post("/extract_coordinates")
async def get_coordinates(req: CoordinateRequest):
    try:
        raw = base64.b64decode(req.file_b64)
        with pdfplumber.open(BytesIO(raw)) as pdf:
            for idx, page in enumerate(pdf.pages):
                text_instances = page.extract_words()
                for word in text_instances:
                    if req.target_text.lower() in word['text'].lower():
                        return {"status": "found", "page": idx, "x": word['x0'], "y": word['top'], "w": word['x1'] - word['x0'], "h": word['bottom'] - word['top']}
        return {"status": "not_found"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- ENDPOINT 8: FIND TEXT COORDS ---
@app.post("/find_text_coords")
async def find_text_coords(req: CoordinateRequest):
    try:
        pdf_bytes = base64.b64decode(req.file_b64)
        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            page_index = len(pdf.pages) - 1
            page = pdf.pages[page_index]
            ph = float(page.height)
            target = req.target_text.upper()
            for word in page.extract_words():
                if target in word['text'].upper():
                    y_coords = ph - float(word['top']) - (float(word['bottom']) - float(word['top']))
                    return {"x": word['x0'], "y": y_coords, "w": 75, "h": 12, "page": page_index, "text_found": word['text']}
        return {"status": "not_found"}
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))


# --- ENDPOINT 9: EDIT INCOTERM ---
@app.post("/edit_incoterm", response_class=Response)
async def edit_incoterm(req: IncotermReq):
    reader = _load_pdf_from_b64(req.file_b64)
    writer = PdfWriter()
    if req.incoterm_change and req.area:
        _overlay_rect_with_text(reader, writer, req.area, req.incoterm_text, font_size=req.font_size, leading=req.leading, debug_outline=req.debug_outline)
    else:
        for p in reader.pages: writer.add_page(p)
    return Response(content=_export(writer), media_type="application/pdf")


# --- ENDPOINT 10: EDIT BILLSHIP ---
@app.post("/edit_billship", response_class=Response)
async def edit_billship(req: BillShipReq):
    reader = _load_pdf_from_b64(req.file_b64)
    writer = PdfWriter()
    bill_area = req.bill_to_area or Rect(page=0, x=72, y=560, w=220, h=80)
    ship_area = req.ship_to_area or Rect(page=0, x=300, y=560, w=250, h=80)
    
    if req.bill_to_text and req.bill_to_area:
        _overlay_rect_with_text(reader, writer, bill_area, req.bill_to_text, font_size=req.font_size, leading=req.leading, debug_outline=req.debug_outline)
    if req.ship_to_text and req.ship_to_area:
        _overlay_rect_with_text(reader, writer, ship_area, req.ship_to_text, font_size=req.font_size, leading=req.leading, debug_outline=req.debug_outline)
        
    if not writer.pages:
        for p in reader.pages: writer.add_page(p)
        
    return Response(content=_export(writer), media_type="application/pdf")


# --- ENDPOINT 11: OVERLAY TEXT BATCH ---
@app.post("/overlay_text_batch", response_class=Response)
async def overlay_text_batch(req: CustomBatchReq):
    reader = _load_pdf_from_b64(req.file_b64)
    if not req.ops:
        w = PdfWriter()
        for p in reader.pages: w.add_page(p)
        return Response(content=_export(w), media_type="application/pdf")
    for op in req.ops:
        w = PdfWriter()
        _overlay_rect_with_text(reader, w, op.area, op.text, font_name=(op.font_name or "Helvetica"), font_size=op.font_size or 9.0, leading=op.leading or 11.0, debug_outline=op.debug_outline or False)
        reader = PdfReader(BytesIO(_export(w)))
    final_writer = PdfWriter()
    for p in reader.pages: final_writer.add_page(p)
    return Response(content=_export(final_writer), media_type="application/pdf")


# --- ENDPOINT 12: CUT RANGE ---
@app.post("/cut_range", response_class=Response)
async def cut_range(req: CutRangeReq):
    reader = _load_pdf_from_b64(req.file_b64)
    writer = PdfWriter()
    total = len(reader.pages)
    start = max(0, min(req.start_page, total - 1))
    end = max(start, min(req.final_page, total - 1))
    for i in range(start, end + 1):
        writer.add_page(reader.pages[i])
    return Response(content=_export(writer), media_type="application/pdf")


# --- ENDPOINT 13: EXTRACT CUSTOM PAGES ---
@app.post("/extract_custom_pages", response_class=Response)
async def extract_custom_pages(req: CustomPagesReq):
    reader = _load_pdf_from_b64(req.file_b64)
    writer = PdfWriter()
    total = len(reader.pages)
    for p_num in req.pages:
        if 0 <= p_num < total:
            writer.add_page(reader.pages[p_num])
    return Response(content=_export(writer), media_type="application/pdf")


# --- ENDPOINT 14: VALIDATE REFERENCE REQUEST ---
@app.post("/validate_reference_request")
def validate_reference_request(req: ReferenceParseReq):
    def normalizar_texto(texto: str) -> str:
        if not texto:
            return ""
        texto_norm = unicodedata.normalize('NFD', texto)
        texto_sin_acentos = ''.join(c for c in texto_norm if unicodedata.category(c) != 'Mn')
        return texto_sin_acentos.upper().strip()

    asunto = normalizar_texto(req.subject)
    cuerpo = normalizar_texto(req.body)
    texto_completo = f"{asunto} {cuerpo}"

    id_tipo = None
    id_valor = None

    cliente_match = re.search(r'\b(CLIENTE|NO\.?\s*CLIENTE|NUMERO\s*DE?\s*CLIENTE)\b.{0,8}?([0-9]{4,6})\b', texto_completo)
    rfc_prefix = re.search(r'\bRFC\s*[:\-\s]\s*([A-Z0-9\-\s]{10,16})\b', texto_completo)
    curp_prefix = re.search(r'\bCURP\s*[:\-\s]\s*([A-Z0-9\-\s]{18,22})\b', texto_completo)

    if cliente_match:
        id_tipo = "NUMERO_CLIENTE"
        id_valor = cliente_match.group(2).strip()
    elif rfc_prefix:
        id_tipo = "RFC"
        id_valor = re.sub(r'[\s\-]', '', rfc_prefix.group(1))
    elif curp_prefix:
        id_tipo = "CURP"
        id_valor = re.sub(r'[\s\-]', '', curp_prefix.group(1))
    else:
        if not curp_prefix:
            curp_prefix = re.search(r'\bCURP\s*[:\-\s]*([A-Z0-9]{18})\b', texto_completo)
        if not rfc_prefix:
            rfc_prefix = re.search(r'\bRFC\s*[:\-\s]*([A-Z0-9]{12,13})\b', texto_completo)

        if rfc_prefix:
            id_tipo = "RFC"
            id_valor = rfc_prefix.group(1).strip()
        elif curp_prefix:
            id_tipo = "CURP"
            id_valor = curp_prefix.group(1).strip()
        else:
            rfc_lenient = re.search(r'\b[A-Z&Ñ]{3,4}[0-9]{6}[A-Z0-9]{3}\b', texto_completo)
            curp_lenient = re.search(r'\b[A-Z&Ñ]{4}[0-9]{6}[A-Z0-9]{8}\b', texto_completo)
            
            if rfc_lenient:
                id_tipo = "RFC"
                id_valor = rfc_lenient.group(0)
            elif curp_lenient:
                id_tipo = "CURP"
                id_valor = curp_lenient.group(0)
            else:
                # ESTADO: RECHAZADO (Sin ID válido)
                return {
                    "status": "rejected",
                    "reason": "No se localizó un identificador legible. Asegúrese de escribir de forma clara su Número de Cliente, RFC o CURP."
                }

    RFC_STRICT = r'^[A-Z&Ñ]{3,4}[0-9]{2}(0[1-9]|1[0-2])(0[1-9]|[12][0-9]|3[01])[A-Z0-9]{3}$'
    CURP_STRICT = r'^[A-Z][AEIOUX][A-Z]{2}[0-9]{2}(0[1-9]|1[0-2])(0[1-9]|[12][0-9]|3[01])[HM][A-Z]{2}[B-DF-HJ-NP-TV-XYZ]{3}[0-9A-Z][0-9]$'
    CLIENTE_STRICT = r'^[0-9]{4,6}$'

    is_valid = False
    if id_tipo == "NUMERO_CLIENTE":
        is_valid = bool(re.match(CLIENTE_STRICT, id_valor))
    elif id_tipo == "RFC":
        is_valid = bool(re.match(RFC_STRICT, id_valor))
    elif id_tipo == "CURP":
        is_valid = bool(re.match(CURP_STRICT, id_valor))

    structure_status = "Válida" if is_valid else "Estructura no válida"

    mapeo_sucursales = {
        "Villas": [r"\bVILLAS?\b", r"\b331\b"],
        "San Miguel de Allende": [r"\bSAN\s+MIGUEL\b", r"\bSAN\s+MIGUEL\s+DE\s+ALLENDE\b", r"\b135\b", r"\bSMA\b"],
        "Guanajuato": [r"\b184\b", r"\bGUANAJUATO\b", r"\bGTO\b"],
        "Solidaridad": [r"\bSOLI\b", r"\bSOLIDARIDAD\b", r"\b192\b"],
        "Silao": [r"\bSILAO\b", r"\b217\b"],
        "Salamanca": [r"\bSALAMANCA\b", r"\b175\b"]
    }

    def obtener_primera_coincidencia(texto: str) -> Optional[str]:
        if not texto:
            return None
        matches = []
        for sucursal, patrones in mapeo_sucursales.items():
            for patron in patrones:
                m = re.search(patron, texto)
                if m:
                    matches.append({"pos": m.start(), "sucursal": sucursal})
                    break
        if not matches:
            return None
        matches.sort(key=lambda x: x["pos"])
        return matches[0]["sucursal"]

    sucursal_encontrada = obtener_primera_coincidencia(asunto)

    if not sucursal_encontrada and cuerpo:
        cuerpo_limpio = re.split(r'_{3,}|={3,}|-{3,}|(?:\r?\n){2,}--\s*', cuerpo)[0]

        for sucursal, patrones in mapeo_sucursales.items():
            for patron in patrones:
                if re.search(r'\b(SUCURSAL|SUC|AGENCIA|PLAZA)\b.{0,15}?' + patron, cuerpo_limpio):
                    sucursal_encontrada = sucursal
                    break
            if sucursal_encontrada:
                break

        if not sucursal_encontrada:
            sucursal_encontrada = obtener_primera_coincidencia(cuerpo_limpio)

    if not sucursal_encontrada and cuerpo:
        sucursal_encontrada = obtener_primera_coincidencia(cuerpo)

    # NUEVO ESTADO: IDCliente (ID válido, pero sin sucursal)
    if not sucursal_encontrada:
        return {
            "status": "IDCliente",
            "search_by": id_tipo,
            "search_value": id_valor,
            "branch": "GENERAL",
            "structure_status": structure_status,
            "reason": "No se introdujo Sucursal, te muestro las referencias que tiene activas el cliente."
        }

    # ESTADO: APROBADO (ID válido y sucursal encontrada)
    return {
        "status": "approved",
        "search_by": id_tipo,
        "search_value": id_valor,
        "branch": sucursal_encontrada,
        "structure_status": structure_status
    }

# --- ENDPOINT 15: DESCARGAR ARCHIVO LOCAL/ALMACENAMIENTO (NUEVO) ---
@app.get("/descargar_archivo/{filepath:path}")
async def descargar_archivo(filepath: str):
    """
    Descarga cualquier archivo almacenado en el microservicio o sus volúmenes.
    Ejemplo de llamada: GET /descargar_archivo/archivos/mi_documento.pdf
    O: GET /descargar_archivo/logo_vegusa.png
    """
    try:
        base_dir = os.path.abspath(os.path.dirname(__file__))
        target_path = os.path.abspath(os.path.join(base_dir, filepath))

        # Validación anti-Directory Traversal (Evita acceder fuera del contenedor/app)
        if not target_path.startswith(base_dir):
            raise HTTPException(
                status_code=400, 
                detail="Acceso denegado: Intento de acceso a una ruta fuera del directorio permitido."
            )

        if not os.path.exists(target_path) or not os.path.isfile(target_path):
            raise HTTPException(
                status_code=404, 
                detail=f"El archivo '{filepath}' no existe o no se encuentra disponible."
            )

        return FileResponse(
            path=target_path, 
            filename=os.path.basename(target_path)
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al descargar archivo: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)